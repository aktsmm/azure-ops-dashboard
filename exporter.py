"""エクスポートユーティリティ（Markdown → Word / PDF）

Markdown レポートを .docx（Word）に変換する。
PDF は docx 経由で Word→PDF 変換（comtypes/LibreOffice）が必要なため、
ここではまず docx 変換を提供する。
"""

from __future__ import annotations

import re
from pathlib import Path
from typing import Any

from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH


def md_to_docx(md_text: str, output_path: Path, title: str = "") -> Path:
    """Markdown テキストを .docx ファイルに変換して保存する。

    Args:
        md_text: Markdown 形式のテキスト
        output_path: 出力先 .docx パス
        title: ドキュメントタイトル（空なら最初の # から取得）

    Returns:
        保存したファイルパス
    """
    doc = Document()

    # スタイル設定
    style = doc.styles["Normal"]
    font = style.font
    font.name = "游ゴシック"
    font.size = Pt(10.5)

    lines = md_text.split("\n")
    i = 0

    # タイトル自動検出
    if not title and lines:
        for line in lines:
            stripped = line.strip()
            if stripped.startswith("# ") and not stripped.startswith("## "):
                title = _strip_md(stripped[2:])
                break

    if title:
        p = doc.add_heading(title, level=0)
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER

    in_code_block = False
    code_lines: list[str] = []
    in_table = False
    table_rows: list[list[str]] = []

    while i < len(lines):
        line = lines[i]
        stripped = line.strip()

        # コードブロック
        if stripped.startswith("```"):
            if in_code_block:
                # コードブロック終了
                _add_code_block(doc, "\n".join(code_lines))
                code_lines = []
                in_code_block = False
            else:
                # テーブル未終了なら閉じる
                if in_table and table_rows:
                    _add_table(doc, table_rows)
                    table_rows = []
                    in_table = False
                in_code_block = True
            i += 1
            continue

        if in_code_block:
            code_lines.append(line)
            i += 1
            continue

        # テーブル
        if stripped.startswith("|") and stripped.endswith("|"):
            cells = [c.strip() for c in stripped.split("|")[1:-1]]
            # セパレータ行（---）はスキップ
            if all(re.match(r"^[-:]+$", c) for c in cells):
                i += 1
                continue
            table_rows.append(cells)
            in_table = True
            i += 1
            continue
        elif in_table and table_rows:
            _add_table(doc, table_rows)
            table_rows = []
            in_table = False

        # 見出し
        if stripped.startswith("#"):
            level = len(stripped) - len(stripped.lstrip("#"))
            level = min(level, 4)
            text = _strip_md(stripped.lstrip("# "))
            if text and not (level == 1 and text == title):
                doc.add_heading(text, level=level)
            i += 1
            continue

        # 空行
        if not stripped:
            i += 1
            continue

        # リスト（箇条書き）
        list_match = re.match(r"^(\s*)([-*+]|\d+\.)\s+(.*)", line)
        if list_match:
            indent = len(list_match.group(1))
            bullet_type = list_match.group(2)
            text = _strip_md(list_match.group(3))
            level_idx = indent // 2
            if re.match(r"\d+\.", bullet_type):
                p = doc.add_paragraph(text, style="List Number")
            else:
                p = doc.add_paragraph(text, style="List Bullet")
            if level_idx > 0:
                p.paragraph_format.left_indent = Inches(0.25 * level_idx)
            i += 1
            continue

        # 水平線
        if stripped in ("---", "***", "___"):
            p = doc.add_paragraph()
            p.add_run("─" * 60).font.color.rgb = RGBColor(0xCC, 0xCC, 0xCC)
            i += 1
            continue

        # 引用
        if stripped.startswith(">"):
            text = _strip_md(stripped.lstrip("> "))
            p = doc.add_paragraph(text)
            p.paragraph_format.left_indent = Inches(0.5)
            if p.runs:
                p.runs[0].font.italic = True
            i += 1
            continue

        # 通常テキスト
        text = _strip_md(stripped)
        if text:
            doc.add_paragraph(text)
        i += 1

    # 未閉じのコードブロック（入力が不正でも欠落しないようベストエフォート）
    if in_code_block and code_lines:
        _add_code_block(doc, "\n".join(code_lines))

    # 未閉じのテーブル
    if in_table and table_rows:
        _add_table(doc, table_rows)

    # 保存
    output_path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(output_path))
    return output_path


def _strip_md(text: str) -> str:
    """Markdown のインライン装飾を除去する。"""
    # Bold + Italic
    text = re.sub(r"\*\*\*(.*?)\*\*\*", r"\1", text)
    # Bold
    text = re.sub(r"\*\*(.*?)\*\*", r"\1", text)
    # Italic
    text = re.sub(r"\*(.*?)\*", r"\1", text)
    # Code
    text = re.sub(r"`(.*?)`", r"\1", text)
    # Link [text](url)
    text = re.sub(r"\[(.*?)\]\(.*?\)", r"\1", text)
    # Emoji shortcuts
    text = re.sub(r":([\w+-]+):", "", text)
    return text.strip()


def _add_code_block(doc: Document, code: str) -> None:
    """コードブロックをグレー背景で追加。"""
    p = doc.add_paragraph()
    run = p.add_run(code)
    run.font.name = "Consolas"
    run.font.size = Pt(9)
    run.font.color.rgb = RGBColor(0xD4, 0xD4, 0xD4)
    p.paragraph_format.left_indent = Inches(0.3)


def _add_table(doc: Document, rows: list[list[str]]) -> None:
    """テーブルをWordテーブルとして追加。"""
    if not rows:
        return
    n_cols = max(len(r) for r in rows)
    table = doc.add_table(rows=len(rows), cols=n_cols, style="Light Grid Accent 1")

    for ri, row in enumerate(rows):
        for ci, cell_text in enumerate(row):
            if ci < n_cols:
                cell = table.cell(ri, ci)
                cell.text = _strip_md(cell_text)
                for paragraph in cell.paragraphs:
                    paragraph.paragraph_format.space_after = Pt(2)
                    for run in paragraph.runs:
                        run.font.size = Pt(9)
    # 1行目をヘッダーとして太字にする
    if rows:
        for ci in range(min(n_cols, len(rows[0]))):
            for run in table.cell(0, ci).paragraphs[0].runs:
                run.bold = True


def md_to_pdf(md_text: str, output_path: Path, title: str = "") -> Path | None:
    """Markdown → PDF 変換。Word経由でPDF化を試みる。

    Windows + Microsoft Word: comtypes 経由
    Mac/Linux: LibreOffice (soffice) 経由
    """
    import sys

    # まず docx を作成
    docx_path = output_path.with_suffix(".docx")
    md_to_docx(md_text, docx_path, title)

    # Windows: comtypes + Microsoft Word
    if sys.platform == "win32":
        try:
            import comtypes.client
            word = None
            doc = None
            try:
                word = comtypes.client.CreateObject("Word.Application")
                word.Visible = False
                doc = word.Documents.Open(str(docx_path.resolve()))
                doc.SaveAs(str(output_path.resolve()), FileFormat=17)  # 17 = wdFormatPDF
                return output_path
            finally:
                try:
                    if doc is not None:
                        doc.Close(False)
                except Exception:
                    pass
                try:
                    if word is not None:
                        word.Quit()
                except Exception:
                    pass
        except Exception:
            pass  # comtypes 不可 → LibreOffice フォールバックへ

    # Mac/Linux (+ Windows fallback): LibreOffice
    try:
        import subprocess
        kwargs: dict[str, Any] = {"capture_output": True, "timeout": 60}
        if sys.platform == "win32":
            kwargs["creationflags"] = subprocess.CREATE_NO_WINDOW
        subprocess.run([
            "soffice", "--headless", "--convert-to", "pdf",
            "--outdir", str(output_path.parent),
            str(docx_path),
        ], **kwargs)
        if output_path.exists():
            return output_path
    except Exception:
        pass
    return None


# ============================================================
# レポート差分比較
# ============================================================

def find_previous_report(output_dir: Path, report_type: str, current_name: str) -> Path | None:
    """output_dir 内で同じ report_type の直前レポートを探す。"""
    pattern = f"{report_type}-report-*.md"
    candidates = sorted(output_dir.glob(pattern), reverse=True)
    for c in candidates:
        if c.name != current_name and c.is_file():
            return c
    return None


def generate_diff_report(prev_path: Path, curr_path: Path) -> str:
    """2つの Markdown レポートの差分を Markdown 形式で返す。

    - セクション（##）単位で追加/削除/変更を検出
    - 行単位の unified diff を付ける
    """
    import difflib

    prev_lines = prev_path.read_text(encoding="utf-8").splitlines(keepends=True)
    curr_lines = curr_path.read_text(encoding="utf-8").splitlines(keepends=True)

    diff = list(difflib.unified_diff(
        prev_lines, curr_lines,
        fromfile=prev_path.name,
        tofile=curr_path.name,
        lineterm="",
    ))

    if not diff:
        return "# 差分レポート\n\n前回と変更はありません。\n"

    # セクション変化サマリ
    prev_sections = _extract_sections(prev_lines)
    curr_sections = _extract_sections(curr_lines)

    added = set(curr_sections) - set(prev_sections)
    removed = set(prev_sections) - set(curr_sections)

    parts: list[str] = []
    parts.append(f"# 差分レポート\n")
    parts.append(f"- 前回: `{prev_path.name}`\n")
    parts.append(f"- 今回: `{curr_path.name}`\n")

    if added:
        parts.append(f"\n## 追加されたセクション\n")
        for s in sorted(added):
            parts.append(f"- {s}\n")

    if removed:
        parts.append(f"\n## 削除されたセクション\n")
        for s in sorted(removed):
            parts.append(f"- {s}\n")

    parts.append(f"\n## 詳細 Diff\n\n```diff\n")
    parts.extend(diff)
    parts.append("\n```\n")

    return "".join(parts)


def _extract_sections(lines: list[str]) -> list[str]:
    """Markdown の ## 見出しを抽出。"""
    sections: list[str] = []
    for line in lines:
        stripped = line.strip() if isinstance(line, str) else ""
        if stripped.startswith("## "):
            sections.append(stripped[3:].strip())
    return sections


# ============================================================
# Markdown バリデーション（レポート品質チェック）
# ============================================================

def validate_markdown(md_text: str) -> list[str]:
    """生成された Markdown レポートを機械的に検証し、警告メッセージのリストを返す。

    チェック項目:
    - 先頭空行
    - テーブルセル内の脚注 [^N]
    - 脚注定義の重複 URL
    - 未定義の脚注参照 / 未使用の脚注定義
    - テーブル列数の不一致
    """
    warnings: list[str] = []
    lines = md_text.split("\n")

    # 1. 先頭空行チェック
    if md_text and md_text[0] in ("\n", "\r", " "):
        warnings.append("先頭に不要な空行/空白があります")

    def _footnote_sort_key(k: str) -> tuple[int, int, str]:
        return (0, int(k), k) if k.isdigit() else (1, 0, k)

    # 2. テーブルセル内の脚注チェック
    footnote_in_cell = re.compile(r"^\|.*\[\^([A-Za-z0-9_-]+)\].*\|")
    for i, line in enumerate(lines, 1):
        if footnote_in_cell.match(line.strip()):
            warnings.append(f"L{i}: テーブルセル内に脚注 [^N] があります（レンダリング崩れの原因）")

    # 3. 脚注定義の収集と重複 URL チェック
    footnote_def = re.compile(r"^\[\^([A-Za-z0-9_-]+)\]:\s*\[.*?\]\((https?://[^\s)]+)\)")
    defined_footnotes: dict[str, str] = {}  # key -> url
    url_to_keys: dict[str, list[str]] = {}
    for line in lines:
        m = footnote_def.match(line.strip())
        if m:
            key, url = m.group(1), m.group(2)
            defined_footnotes[key] = url
            url_to_keys.setdefault(url, []).append(key)

    for url, keys in url_to_keys.items():
        if len(keys) > 1:
            warnings.append(f"脚注 [{', '.join(keys)}] が同一 URL を重複定義しています: {url[:80]}")

    # 4. 脚注参照 vs 定義の整合性
    ref_pattern = re.compile(r"\[\^([A-Za-z0-9_-]+)\]")
    referenced: set[str] = set()
    for line in lines:
        if not line.strip().startswith("[^"):
            # 本文中の参照
            referenced.update(ref_pattern.findall(line))

    defined_set = set(defined_footnotes.keys())
    undefined = referenced - defined_set
    unused = defined_set - referenced
    if undefined:
        warnings.append(f"未定義の脚注参照: {', '.join(sorted(undefined, key=_footnote_sort_key))}")
    if unused:
        warnings.append(f"未使用の脚注定義: {', '.join(sorted(unused, key=_footnote_sort_key))}")

    # 5. テーブル列数の一貫性チェック
    in_table = False
    table_col_count = 0
    table_start_line = 0
    for i, line in enumerate(lines, 1):
        stripped = line.strip()
        if stripped.startswith("|") and stripped.endswith("|"):
            col_count = stripped.count("|") - 1
            if not in_table:
                in_table = True
                table_col_count = col_count
                table_start_line = i
            else:
                if col_count != table_col_count:
                    warnings.append(
                        f"L{i}: テーブル列数が不一致（ヘッダー={table_col_count}, この行={col_count}、開始L{table_start_line}）"
                    )
        else:
            in_table = False

    return warnings


def remove_unused_footnote_definitions(md_text: str) -> tuple[str, list[str]]:
    """未使用の脚注定義を削除する（ベストエフォート）。

    - 参照されていない `[^N]: ...` 定義ブロックを除去する。
    - 脚注定義が複数行（次行以降がインデント）でもまとめて除去する。

    Returns:
        (cleaned_markdown, removed_keys)
    """
    lines = md_text.split("\n")
    if not lines:
        return md_text, []

    def _footnote_sort_key(k: str) -> tuple[int, int, str]:
        return (0, int(k), k) if k.isdigit() else (1, 0, k)

    ref_pattern = re.compile(r"\[\^([A-Za-z0-9_-]+)\]")
    def_pattern = re.compile(r"^\[\^([A-Za-z0-9_-]+)\]:")

    referenced: set[str] = set()
    for line in lines:
        if not line.strip().startswith("[^"):
            referenced.update(ref_pattern.findall(line))

    def_starts: dict[int, str] = {}
    for idx, line in enumerate(lines):
        m = def_pattern.match(line.strip())
        if m:
            def_starts[idx] = m.group(1)

    if not def_starts:
        return md_text, []

    remove_line: set[int] = set()
    removed_keys: set[str] = set()

    sorted_starts = sorted(def_starts.items())
    for pos, (start_idx, key) in enumerate(sorted_starts):
        if key in referenced:
            continue

        # Remove the definition line and its continuation lines.
        # Continuation lines are indented (2+ spaces or a tab) and are not another footnote definition.
        end_idx = sorted_starts[pos + 1][0] if pos + 1 < len(sorted_starts) else len(lines)
        remove_line.add(start_idx)
        for j in range(start_idx + 1, end_idx):
            nxt = lines[j]
            if def_pattern.match(nxt.strip()):
                break
            if nxt.startswith("\t") or nxt.startswith("  "):
                remove_line.add(j)
                continue
            break

        removed_keys.add(key)

    if not removed_keys:
        return md_text, []

    cleaned = "\n".join(line for i, line in enumerate(lines) if i not in remove_line).rstrip() + "\n"
    return cleaned, sorted(removed_keys, key=_footnote_sort_key)
