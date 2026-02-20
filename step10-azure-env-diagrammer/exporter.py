"""Step10: エクスポートユーティリティ（Markdown → Word / PDF）

Markdown レポートを .docx（Word）に変換する。
PDF は docx 経由で Word→PDF 変換（comtypes/LibreOffice）が必要なため、
ここではまず docx 変換を提供する。
"""

from __future__ import annotations

import re
from pathlib import Path
from typing import Any

from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH


def md_to_docx(md_text: str, output_path: Path, title: str = "") -> Path:
    """Markdown テキストを .docx ファイルに変換して保存する。

    Args:
        md_text: Markdown 形式のテキスト
        output_path: 出力先 .docx パス
        title: ドキュメントタイトル（空なら最初の # から取得）

    Returns:
        保存したファイルパス
    """
    doc = Document()

    # スタイル設定
    style = doc.styles["Normal"]
    font = style.font
    font.name = "游ゴシック"
    font.size = Pt(10.5)

    lines = md_text.split("\n")
    i = 0

    # タイトル自動検出
    if not title and lines:
        for line in lines:
            stripped = line.strip()
            if stripped.startswith("# ") and not stripped.startswith("## "):
                title = _strip_md(stripped[2:])
                break

    if title:
        p = doc.add_heading(title, level=0)
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER

    in_code_block = False
    code_lines: list[str] = []
    in_table = False
    table_rows: list[list[str]] = []

    while i < len(lines):
        line = lines[i]
        stripped = line.strip()

        # コードブロック
        if stripped.startswith("```"):
            if in_code_block:
                # コードブロック終了
                _add_code_block(doc, "\n".join(code_lines))
                code_lines = []
                in_code_block = False
            else:
                # テーブル未終了なら閉じる
                if in_table and table_rows:
                    _add_table(doc, table_rows)
                    table_rows = []
                    in_table = False
                in_code_block = True
            i += 1
            continue

        if in_code_block:
            code_lines.append(line)
            i += 1
            continue

        # テーブル
        if stripped.startswith("|") and stripped.endswith("|"):
            cells = [c.strip() for c in stripped.split("|")[1:-1]]
            # セパレータ行（---）はスキップ
            if all(re.match(r"^[-:]+$", c) for c in cells):
                i += 1
                continue
            table_rows.append(cells)
            in_table = True
            i += 1
            continue
        elif in_table and table_rows:
            _add_table(doc, table_rows)
            table_rows = []
            in_table = False

        # 見出し
        if stripped.startswith("#"):
            level = len(stripped) - len(stripped.lstrip("#"))
            level = min(level, 4)
            text = _strip_md(stripped.lstrip("# "))
            if text and not (level == 1 and text == title):
                doc.add_heading(text, level=level)
            i += 1
            continue

        # 空行
        if not stripped:
            i += 1
            continue

        # リスト（箇条書き）
        list_match = re.match(r"^(\s*)([-*+]|\d+\.)\s+(.*)", line)
        if list_match:
            indent = len(list_match.group(1))
            bullet_type = list_match.group(2)
            text = _strip_md(list_match.group(3))
            level_idx = indent // 2
            if re.match(r"\d+\.", bullet_type):
                p = doc.add_paragraph(text, style="List Number")
            else:
                p = doc.add_paragraph(text, style="List Bullet")
            if level_idx > 0:
                p.paragraph_format.left_indent = Inches(0.25 * level_idx)
            i += 1
            continue

        # 水平線
        if stripped in ("---", "***", "___"):
            p = doc.add_paragraph()
            p.add_run("─" * 60).font.color.rgb = RGBColor(0xCC, 0xCC, 0xCC)
            i += 1
            continue

        # 引用
        if stripped.startswith(">"):
            text = _strip_md(stripped.lstrip("> "))
            p = doc.add_paragraph(text)
            p.paragraph_format.left_indent = Inches(0.5)
            p.runs[0].font.italic = True if p.runs else None
            i += 1
            continue

        # 通常テキスト
        text = _strip_md(stripped)
        if text:
            doc.add_paragraph(text)
        i += 1

    # 未閉じのテーブル
    if in_table and table_rows:
        _add_table(doc, table_rows)

    # 保存
    output_path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(output_path))
    return output_path


def _strip_md(text: str) -> str:
    """Markdown のインライン装飾を除去する。"""
    # Bold + Italic
    text = re.sub(r"\*\*\*(.*?)\*\*\*", r"\1", text)
    # Bold
    text = re.sub(r"\*\*(.*?)\*\*", r"\1", text)
    # Italic
    text = re.sub(r"\*(.*?)\*", r"\1", text)
    # Code
    text = re.sub(r"`(.*?)`", r"\1", text)
    # Link [text](url)
    text = re.sub(r"\[(.*?)\]\(.*?\)", r"\1", text)
    # Emoji shortcuts
    text = re.sub(r":([\w+-]+):", "", text)
    return text.strip()


def _add_code_block(doc: Document, code: str) -> None:
    """コードブロックをグレー背景で追加。"""
    p = doc.add_paragraph()
    run = p.add_run(code)
    run.font.name = "Consolas"
    run.font.size = Pt(9)
    run.font.color.rgb = RGBColor(0xD4, 0xD4, 0xD4)
    p.paragraph_format.left_indent = Inches(0.3)


def _add_table(doc: Document, rows: list[list[str]]) -> None:
    """テーブルをWordテーブルとして追加。"""
    if not rows:
        return
    n_cols = max(len(r) for r in rows)
    table = doc.add_table(rows=len(rows), cols=n_cols, style="Light Grid Accent 1")

    for ri, row in enumerate(rows):
        for ci, cell_text in enumerate(row):
            if ci < n_cols:
                cell = table.cell(ri, ci)
                cell.text = _strip_md(cell_text)
                for paragraph in cell.paragraphs:
                    paragraph.paragraph_format.space_after = Pt(2)
                    for run in paragraph.runs:
                        run.font.size = Pt(9)
    # 1行目をヘッダーとして太字にする
    if rows:
        for ci in range(min(n_cols, len(rows[0]))):
            for run in table.cell(0, ci).paragraphs[0].runs:
                run.bold = True


def md_to_pdf(md_text: str, output_path: Path, title: str = "") -> Path | None:
    """Markdown → PDF 変換。Word経由でPDF化を試みる。

    Windows + Microsoft Word インストール済みの場合のみ動作。
    """
    import sys
    if sys.platform != "win32":
        return None

    # まず docx を作成
    docx_path = output_path.with_suffix(".docx")
    md_to_docx(md_text, docx_path, title)

    try:
        import comtypes.client
        word = None
        doc = None
        try:
            word = comtypes.client.CreateObject("Word.Application")
            word.Visible = False
            doc = word.Documents.Open(str(docx_path.resolve()))
            doc.SaveAs(str(output_path.resolve()), FileFormat=17)  # 17 = wdFormatPDF
            return output_path
        finally:
            try:
                if doc is not None:
                    doc.Close(False)
            except Exception:
                pass
            try:
                if word is not None:
                    word.Quit()
            except Exception:
                pass
    except Exception:
        # comtypes がなければ LibreOffice を試す
        try:
            import subprocess
            subprocess.run([
                "soffice", "--headless", "--convert-to", "pdf",
                "--outdir", str(output_path.parent),
                str(docx_path),
            ], capture_output=True, timeout=60)
            if output_path.exists():
                return output_path
        except Exception:
            pass
    return None
